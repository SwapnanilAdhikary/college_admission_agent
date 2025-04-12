from docx import Document

# Create a new Document object
doc = Document()

# Add the Student Resume section
doc.add_heading("Student Resume", level=1)
doc.add_paragraph("Name: Ram Sharma")
doc.add_paragraph("Date of Birth: 2005-06-15")
doc.add_paragraph("Age at Admission: 17 (less than 21)")
doc.add_paragraph("Address: 123 Academic Lane, Study City, Country")
doc.add_paragraph("Contact: +91-9876543210")
doc.add_paragraph("Email: ram.sharma@example.com")

# Add the Academic Qualification and Marksheet section
doc.add_heading("Academic Qualification and Marksheet", level=1)
doc.add_paragraph("Institution: Central Higher Secondary School")
doc.add_paragraph("Board: XYZ Higher Secondary Board")
doc.add_paragraph("Year of Passing: 2023")
doc.add_paragraph("Subjects: Physics, Chemistry, Mathematics, English, and others")
doc.add_paragraph("Marks in PCM:")
doc.add_paragraph("  - Physics: 95%")
doc.add_paragraph("  - Chemistry: 93%")
doc.add_paragraph("  - Mathematics: 92%")
doc.add_paragraph("Aggregate in PCM Subjects: 93%")
doc.add_paragraph("Overall Percentage: 91%")
doc.add_paragraph("Remarks: Passed with distinction")

# Add the Parent Income Certificate section
doc.add_heading("Parent Income Certificate", level=1)
doc.add_paragraph(
    "This is to certify that the annual income of the parents of Ram Sharma is ₹200,000. "
    "All details are verified by the issuing authority."
)

# Add the Loan Request Document section
doc.add_heading("Loan Request Document", level=1)
doc.add_paragraph(
    "Ram Sharma hereby requests a student loan to support his higher education expenses. "
    "Reason for Loan: To cover tuition fees and related educational expenses."
)
doc.add_paragraph(
    "I, the undersigned, affirm that the above income details are correct and request the issuance of a loan "
    "per the university’s loan policy (loan will be approved only if the annual parent income is less than ₹2,50,000)."
)

# Save the document
doc.save("Ram_Sharma_Application.docx")
print("Document 'Ram_Sharma_Application.docx' created successfully.")
